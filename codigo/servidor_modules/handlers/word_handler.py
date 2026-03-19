# servidor_modules/handlers/word_handler.py
"""
word_handler.py - Manipulação de documentos Word
"""

import json
import os
import tempfile
from pathlib import Path
from datetime import datetime
from docxtpl import DocxTemplate
import traceback
from http.server import BaseHTTPRequestHandler
from typing import Dict, List, Any, Optional

class WordHandler:
    """Handler para geração de documentos Word"""
    
    def __init__(self, project_root, file_utils):
        self.project_root = project_root
        self.file_utils = file_utils
        self.templates_dir = project_root / "word_templates"
        self.ensure_templates_dir()
        
    def ensure_templates_dir(self):
        """Garante que a pasta de templates existe"""
        self.templates_dir.mkdir(exist_ok=True)
        
        # Cria templates padrão se não existirem
        default_templates = {
            "proposta_comercial_template.docx": {
                "name": "Proposta Comercial",
                "description": "Documento comercial com valores, condições de pagamento"
            },
            "proposta_tecnica_template.docx": {
                "name": "Proposta Técnica", 
                "description": "Documento técnico com especificações e cálculos"
            }
        }
        
        # Cria arquivos de placeholder se não existirem
        for filename, info in default_templates.items():
            template_path = self.templates_dir / filename
            if not template_path.exists():
                self.create_placeholder_template(template_path, info["name"])
                
    def create_placeholder_template(self, template_path, template_name):
        """Cria um template placeholder se não existir"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Título
            title = doc.add_heading(f'Template: {template_name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Informações
            doc.add_paragraph(f'Template criado em: {datetime.now().strftime("%d/%m/%Y")}')
            doc.add_paragraph('Este é um template placeholder. Substitua com seu template real.')
            doc.add_paragraph('Variáveis disponíveis:')
            
            # Variáveis exemplo para Proposta Comercial
            if "comercial" in template_name.lower():
                vars_list = doc.add_paragraph()
                vars_list.add_run('• {{data_emissao}} - Data de emissão\n').bold = True
                vars_list.add_run('• {{empresa_nome}} - Nome da empresa cliente\n')
                vars_list.add_run('• {{obra_nome}} - Nome da obra\n')
                vars_list.add_run('• {{cliente_final}} - Nome do cliente final\n')
                vars_list.add_run('• {{valor_total_projeto}} - Valor total do projeto\n')
                vars_list.add_run('• {{total_global}} - Valor total global\n')
                vars_list.add_run('• {{aplicacoes_groups}} - Lista de aplicações com máquinas, dutos e acessórios\n')
                vars_list.add_run('• {{engenharia_valor}} - Valor da engenharia\n')
                vars_list.add_run('• {{engenharia_descricao}} - Descrição da engenharia\n')
                vars_list.add_run('• {{adicionais}} - Lista de serviços adicionais\n')
            # Variáveis exemplo para Proposta Técnica
            elif "tecnica" in template_name.lower():
                vars_list = doc.add_paragraph()
                vars_list.add_run('• {{data_emissao}} - Data de emissão\n').bold = True
                vars_list.add_run('• {{empresa_nome}} - Nome da empresa\n')
                vars_list.add_run('• {{projetos}} - Lista de projetos com salas e máquinas\n')
                vars_list.add_run('• {{opcoes_por_tipo}} - Opções disponíveis por tipo de máquina\n')
                vars_list.add_run('• {{tensoes_disponiveis}} - Tensões elétricas disponíveis\n')
            
            doc.save(str(template_path))
            print(f"✅ Template placeholder criado: {template_path}")
            return True
        except Exception as e:
            print(f"❌ Erro ao criar template placeholder: {e}")
            return False
    
    def get_available_templates(self):
        """Retorna templates disponíveis"""
        templates = []
        for file in self.templates_dir.glob("*.docx"):
            templates.append({
                "filename": file.name,
                "path": str(file),
                "size": file.stat().st_size,
                "modified": datetime.fromtimestamp(file.stat().st_mtime).isoformat()
            })
        return templates
    
    def get_obra_data(self, obra_id):
        """Obtém dados completos de uma obra"""
        try:
            backup_file = self.project_root / "json" / "backup.json"
            if not backup_file.exists():
                return None
                
            with open(backup_file, "r", encoding="utf-8") as f:
                backup_data = json.load(f)
            
            obras = backup_data.get("obras", [])
            for obra in obras:
                if str(obra.get("id")) == obra_id:
                    return obra
            return None
        except Exception as e:
            print(f"❌ Erro ao buscar obra: {e}")
            return None
    
    def generate_context_for_pc(self, obra_id: str) -> Optional[Dict]:
        """Gera contexto para Proposta Comercial - MÉTODO LEGADO"""
        try:
            # Importar o gerador avançado
            from servidor_modules.generators.wordPC_generator import WordPCGenerator
            
            # Criar instância do gerador
            pc_generator = WordPCGenerator(self.project_root, self.file_utils)
            
            # Usar o método do gerador
            return pc_generator.generate_context_for_pc(obra_id)
                
        except Exception as e:
            print(f"❌ Erro em generate_context_for_pc: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def generate_proposta_comercial(self, obra_id: str, template_path: Path) -> Optional[str]:
        """Gera documento de Proposta Comercial com tratamento de erros melhorado"""
        try:
            # Verificar template
            if not template_path.exists():
                print(f"❌ Template não encontrado: {template_path}")
                return None
            
            # Verificar se é um arquivo válido
            if template_path.stat().st_size == 0:
                print(f"❌ Template está vazio: {template_path}")
                return None
            
            # Usar o gerador avançado
            from servidor_modules.generators.wordPC_generator import WordPCGenerator
            pc_generator = WordPCGenerator(self.project_root, self.file_utils)
            
            # Gerar contexto
            context = pc_generator.generate_context_for_pc(obra_id)
            if not context:
                raise ValueError("Não foi possível gerar contexto para a PC")
            
            print(f" Contexto gerado com {len(context.get('aplicacoes_groups', []))} aplicações")
            
            # Testar contexto básico primeiro
            test_context = {
                "data_emissao": context.get("data_emissao", ""),
                "empresa_nome": context.get("empresa_nome", ""),
                "obra_nome": context.get("obra_nome", ""),
                "cliente_final": context.get("cliente_final", ""),
                "projeto_nome": context.get("projeto_nome", ""),
                "aplicacoes_groups": [],
                "engenharia_valor": context.get("engenharia_valor", ""),
                "engenharia_descricao": context.get("engenharia_descricao", ""),
                "tem_adicionais": False,
                "adicionais": [],
                "valor_total_projeto": context.get("valor_total_projeto", ""),
                "total_global": context.get("total_global", "")
            }
            
            print("🧪 Testando template com contexto básico...")
            
            try:
                doc = DocxTemplate(str(template_path))
                doc.render(test_context)
                print("✅ Template básico funciona!")
            except Exception as template_error:
                print(f"❌ Erro no template: {template_error}")
            
            # Agora renderizar com contexto completo
            print(" Renderizando com contexto completo...")
            doc = DocxTemplate(str(template_path))
            doc.render(context)
            
            # Salvar arquivo temporário
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                output_path = tmp.name
                doc.save(output_path)
            
            print(f"✅ Proposta Comercial gerada: {output_path}")
            return output_path
            
        except Exception as e:
            print(f"❌ Erro ao gerar Proposta Comercial: {e}")
            traceback.print_exc()
            return None
    
    # ===============================================================
    # NOVO MÉTODO: GERAÇÃO DE PROPOSTA TÉCNICA USANDO GERADOR AVANÇADO
    # ===============================================================
    def generate_proposta_tecnica_avancada(self, obra_id):
        """Gera proposta técnica usando o gerador avançado WordPTGenerator"""
        try:
            # Importar o gerador avançado para PT
            from servidor_modules.generators.wordPT_generator import WordPTGenerator
            
            # Criar instância do gerador
            pt_generator = WordPTGenerator(self.project_root, self.file_utils)
            
            # Localizar template
            template_path = self.templates_dir / "proposta_tecnica_template.docx"
            if not template_path.exists():
                # Tentar encontrar qualquer template .docx
                docx_files = list(self.templates_dir.glob("*.docx"))
                if docx_files:
                    template_path = docx_files[0]
                else:
                    return None, None, "Nenhum template encontrado na pasta word_templates"
            
            # Validar obra (opcional, não bloqueia)
            is_valid, message = self.validate_obra_for_pc(obra_id)  # reutiliza validação
            if not is_valid:
                print(f"⚠️ Validação: {message}")
            
            # Gerar documento
            output_path = pt_generator.generate_proposta_tecnica(obra_id, template_path)
            
            if output_path:
                # Gerar nome do arquivo usando o método do handler
                obra_data = self.get_obra_data(obra_id)
                if obra_data:
                    filename = self.generate_pt_filename(obra_data)
                else:
                    # Fallback
                    from datetime import datetime
                    import pytz
                    try:
                        tz = pytz.timezone('America/Sao_Paulo')
                        data_atual = datetime.now(tz)
                    except:
                        data_atual = datetime.now()
                    filename = f"PT_Obra_{obra_id}_{data_atual.strftime('%d-%m-%Y')}.docx"
                
                # Log de geração
                obra_summary = self.get_obra_summary(obra_id)
                print(f"📄 Proposta Técnica gerada para obra: {obra_summary.get('obra_nome')}")
                print(f"   - Arquivo: {filename}")
                
                return output_path, filename, None
            else:
                return None, None, "Falha ao gerar documento"
                
        except Exception as e:
            print(f"❌ Erro em generate_proposta_tecnica_avancada: {e}")
            import traceback
            traceback.print_exc()
            return None, None, str(e)
    
    def generate_context_for_obra(self, obra_data, template_type="comercial"):
        """Gera contexto para preenchimento do template (método genérico - mantido para compatibilidade)"""
        try:
            # Dados básicos da obra
            obra_nome = obra_data.get("nome", "Obra não especificada")
            cliente = obra_data.get("cliente", {})
            cliente_nome = cliente.get("nome", "Cliente não especificado") if isinstance(cliente, dict) else "Cliente não especificado"
            
            # Endereço
            endereco_completo = ""
            if isinstance(cliente, dict):
                endereco_parts = []
                if cliente.get("endereco"):
                    endereco_parts.append(cliente["endereco"])
                if cliente.get("bairro"):
                    endereco_parts.append(cliente["bairro"])
                if cliente.get("cidade"):
                    endereco_parts.append(cliente["cidade"])
                if cliente.get("estado"):
                    endereco_parts.append(cliente["estado"])
                if cliente.get("cep"):
                    endereco_parts.append(f"CEP: {cliente['cep']}")
                endereco_completo = ", ".join(filter(None, endereco_parts))
            
            # Formatar valores
            def formatar_valor(valor):
                return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            
            # Data atual - usar data de geração, não data de cadastro
            from datetime import datetime
            import pytz
            
            try:
                tz = pytz.timezone('America/Sao_Paulo')
                data_atual = datetime.now(tz)
            except:
                data_atual = datetime.now()
            
            # Contexto base
            context = {
                "obra_nome": obra_nome,
                "cliente_nome": cliente_nome,
                "endereco": endereco_completo,
                "data_emissao": data_atual.strftime("%d/%m/%Y"),  # Data atual
                "data_emissao_completa": data_atual.strftime("%d de %B de %Y"),
                "hora_emissao": data_atual.strftime("%H:%M"),
            }
            
            # Adicionar dados específicos por tipo de template
            if template_type == "comercial":
                context.update({
                    "titulo_documento": "PROPOSTA COMERCIAL",
                    "tipo_proposta": "Comercial",
                    "condicoes_pagamento": "50% na assinatura do contrato, 50% na entrega",
                    "validade_proposta": "10 dias úteis",
                    "garantia": "12 meses",
                    "prazo_entrega": "60 dias úteis",
                })
            elif template_type == "tecnica":
                context.update({
                    "titulo_documento": "PROPOSTA TÉCNICA",
                    "tipo_proposta": "Técnica",
                    "normas_aplicaveis": "NBR 16401, NBR 7256, NBR 14606",
                    "escopo_trabalho": "Fornecimento e instalação completa do sistema de climatização",
                    "memoria_calculo": "Cálculos realizados conforme normas técnicas vigentes",
                    "especificacoes_tecnicas": "Todos os equipamentos conforme catálogo técnico",
                })
            
            return context
        except Exception as e:
            print(f"❌ Erro ao gerar contexto: {e}")
            return {}
        
    def generate_proposta_comercial_avancada(self, obra_id):
        """Gera proposta comercial usando o gerador avançado - VERSÃO CORRIGIDA"""
        try:
            # Importar o gerador avançado
            from servidor_modules.generators.wordPC_generator import WordPCGenerator
            
            # Criar instância do gerador
            pc_generator = WordPCGenerator(self.project_root, self.file_utils)
            
            # Localizar template
            template_path = self.templates_dir / "proposta_comercial_template.docx"
            if not template_path.exists():
                # Tentar encontrar qualquer template .docx
                docx_files = list(self.templates_dir.glob("*.docx"))
                if docx_files:
                    template_path = docx_files[0]
                else:
                    return None, None, "Nenhum template encontrado na pasta word_templates"
            
            # Validar obra (AGORA NÃO BLOQUEIA POR FALTA DE ITENS)
            is_valid, message = self.validate_obra_for_pc(obra_id)
            if not is_valid:
                # Só retorna erro se for algo crítico (obra não existe, campos obrigatórios faltando)
                return None, None, message
            
            # Gerar proposta (AGORA FUNCIONA MESMO SEM ITENS)
            output_path = pc_generator.generate_proposta_comercial(obra_id, template_path)
            
            if output_path:
                # Gerar nome do arquivo usando o método do gerador
                obra_data = self.get_obra_data(obra_id)
                if obra_data:
                    filename = pc_generator.generate_filename(obra_data, "comercial")
                else:
                    # Fallback
                    from datetime import datetime
                    import pytz
                    try:
                        tz = pytz.timezone('America/Sao_Paulo')
                        data_atual = datetime.now(tz)
                    except:
                        data_atual = datetime.now()
                    filename = f"PC_Obra_{obra_id}_{data_atual.strftime('%d-%m-%Y')}.docx"
                
                # Log de geração
                obra_summary = self.get_obra_summary(obra_id)
                print(f"📄 Proposta Comercial gerada para obra: {obra_summary.get('obra_nome')}")
                print(f"   - Total: {obra_summary.get('valor_total_formatado')}")
                print(f"   - Máquinas: {obra_summary.get('total_machines')}")
                print(f"   - Arquivo: {filename}")
                
                return output_path, filename, None
            else:
                return None, None, "Falha ao gerar documento"
                
        except Exception as e:
            print(f"❌ Erro em generate_proposta_comercial_avancada: {e}")
            import traceback
            traceback.print_exc()
            return None, None, str(e)
            
    def generate_word_document(self, obra_id, template_type="comercial"):
        """Gera documento Word baseado no template"""
        try:
            # Para Proposta Comercial, usar o método avançado
            if template_type == "comercial":
                return self.generate_proposta_comercial_avancada(obra_id)
            # Para Proposta Técnica, usar o método avançado
            elif template_type == "tecnica":
                return self.generate_proposta_tecnica_avancada(obra_id)
            else:
                return None, None, f"Tipo de template não suportado: {template_type}"
                    
        except Exception as e:
            print(f"❌ Erro na geração do Word: {e}")
            import traceback
            traceback.print_exc()
            return None, None, str(e)
    
    def generate_both_documents(self, obra_id):
        """Gera ambos os documentos (comercial e técnico)"""
        try:
            # Gerar proposta comercial
            pc_path, pc_filename, pc_error = self.generate_proposta_comercial_avancada(obra_id)
            if pc_error:
                return None, None, pc_error
            
            # Gerar proposta técnica
            pt_path, pt_filename, pt_error = self.generate_proposta_tecnica_avancada(obra_id)
            if pt_error:
                # Limpar arquivo gerado anteriormente
                if pc_path and os.path.exists(pc_path):
                    os.unlink(pc_path)
                return None, None, pt_error
            
            # Para ambos, criar um ZIP com os dois arquivos
            import zipfile
            import tempfile
            from datetime import datetime
            
            # Criar arquivo ZIP temporário
            with tempfile.NamedTemporaryFile(suffix='.zip', delete=False) as tmp_zip:
                zip_path = tmp_zip.name
            
            # Criar ZIP com os dois documentos
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                # Adicionar PC
                if pc_path and os.path.exists(pc_path):
                    zipf.write(pc_path, pc_filename)
                
                # Adicionar PT
                if pt_path and os.path.exists(pt_path):
                    zipf.write(pt_path, pt_filename)
            
            # Limpar arquivos individuais
            if pc_path and os.path.exists(pc_path):
                os.unlink(pc_path)
            if pt_path and os.path.exists(pt_path):
                os.unlink(pt_path)
            
            # Gerar nome do arquivo ZIP
            obra_data = self.get_obra_data(obra_id)
            if obra_data:
                # Extrair base do nome (sigla_numero)
                base_name = self.extract_filename_base(obra_data)
                zip_filename = f"PC_PT_{base_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
            else:
                zip_filename = f"PC_PT_{obra_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
            
            return zip_path, zip_filename, None
            
        except Exception as e:
            print(f"❌ Erro ao gerar ambos documentos: {e}")
            return None, None, str(e)
    
    def get_machine_types_with_specifications(self):
        """Obtém tipos de máquinas com suas especificações do BD"""
        try:
            dados_file = self.project_root / "json" / "dados.json"
            if not dados_file.exists():
                return []
            
            with open(dados_file, "r", encoding="utf-8") as f:
                dados_data = json.load(f)
            
            machines = dados_data.get("machines", [])
            machine_types = []
            
            for machine in machines:
                machine_type = machine.get("type", "")
                especificacao = machine.get("especificacao", "")
                
                if machine_type:
                    machine_types.append({
                        "type": machine_type,
                        "especificacao": especificacao if especificacao else "Não especificada",
                        "has_impostos": "impostos" in machine,
                        "has_options": "options" in machine and machine["options"]
                    })
            
            return machine_types
            
        except Exception as e:
            print(f"❌ Erro ao obter tipos de máquinas: {e}")
            return []
    
    def validate_obra_for_pc(self, obra_id):
        """Valida se a obra tem todos os dados necessários para gerar PC"""
        try:
            obra_data = self.get_obra_data(obra_id)
            if not obra_data:
                print(f"⚠️ Obra {obra_id} não encontrada - documento será gerado com valores padrão")
                return True, "Obra não encontrada, usando valores padrão"
            
            obra_nome = obra_data.get("nome")
            empresa_nome = obra_data.get("empresaNome")
            cliente_final = obra_data.get("clienteFinal")
            
            if not obra_nome:
                print("⚠️ Campo 'nome' não encontrado - usando valor padrão")
            if not empresa_nome:
                print("⚠️ Campo 'empresaNome' não encontrado - usando valor padrão")
            if not cliente_final:
                print("⚠️ Campo 'clienteFinal' não encontrado - usando valor padrão")
            
            # Projetos (pode ser vazio)
            projetos = obra_data.get("projetos", [])
            if not projetos:
                print("⚠️ Obra não tem projetos - documento será gerado sem projetos")
            
            # Itens (pode ser vazio) - apenas verificar para log
            has_items = False
            for projeto in projetos:
                if isinstance(projeto, dict):
                    salas = projeto.get("salas", [])
                    for sala in salas:
                        if isinstance(sala, dict):
                            if (sala.get("maquinas") or 
                                sala.get("dutos") or 
                                sala.get("acessorios")):
                                has_items = True
                                break
                if has_items:
                    break
            
            if not has_items:
                print("⚠️ Nenhum item encontrado - documento será gerado sem itens")
            
            return True, "Documento será gerado com os dados disponíveis"
            
        except Exception as e:
            print(f"❌ Erro ao validar obra: {e}")
            # ⚠️ NEM MESMO ERRO DEVE BLOQUEAR - APENAS LOGAR
            return True, f"Erro na validação, mas documento será gerado: {str(e)}"
    
    def get_obra_summary(self, obra_id):
        """Retorna resumo da obra para debug/log"""
        try:
            obra_data = self.get_obra_data(obra_id)
            if not obra_data:
                return {"error": "Obra não encontrada"}
            
            projetos = obra_data.get("projetos", [])
            total_machines = 0
            total_dutos = 0
            total_acessorios = 0
            total_value = obra_data.get("valorTotalObra", 0)
            
            for projeto in projetos:
                if isinstance(projeto, dict):
                    salas = projeto.get("salas", [])
                    for sala in salas:
                        if isinstance(sala, dict):
                            maquinas = sala.get("maquinas", [])
                            dutos = sala.get("dutos", [])
                            acessorios = sala.get("acessorios", [])
                            total_machines += len(maquinas)
                            total_dutos += len(dutos)
                            total_acessorios += len(acessorios)
            
            return {
                "obra_id": obra_id,
                "obra_nome": obra_data.get("nome", ""),
                "empresa_nome": obra_data.get("empresaNome", ""),
                "cliente_final": obra_data.get("clienteFinal", ""),
                "numero_projetos": len(projetos),
                "total_machines": total_machines,
                "total_dutos": total_dutos,
                "total_acessorios": total_acessorios,
                "valor_total": total_value,
                "valor_total_formatado": f"R$ {total_value:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            }
            
        except Exception as e:
            print(f"❌ Erro ao obter resumo da obra: {e}")
            return {"error": str(e)}
        
    def extract_filename_base(self, obra_data):
        """Extrai base do nome do arquivo (sigla_numero)"""
        try:
            import re
            
            # Extrair sigla
            sigla = obra_data.get("empresaSigla", "")
            if not sigla:
                empresa_nome = obra_data.get("empresaNome", "")
                if empresa_nome:
                    # Extrair sigla entre parênteses
                    match = re.search(r'\(([^)]+)\)', empresa_nome)
                    if match:
                        sigla = match.group(1)
                    else:
                        # Usar iniciais (primeiras 3 palavras, máximo 5 caracteres)
                        palavras = empresa_nome.split()
                        if palavras:
                            # Pegar iniciais das primeiras 3 palavras
                            iniciais = ''.join([p[0].upper() for p in palavras[:3] if p and p[0].isalpha()])
                            sigla = iniciais[:5] if iniciais else "EMP"
            
            if not sigla:
                sigla = "EMP"
            
            # Extrair número do cliente
            cliente_numero = obra_data.get("clienteNumero", "")
            if not cliente_numero:
                # Verificar se tem número no nome do cliente
                cliente_final = obra_data.get("clienteFinal", "")
                if cliente_final:
                    # Procurar por padrões como "Cliente 001" ou "001 - Cliente"
                    numeros = re.findall(r'\b(\d{2,})\b', cliente_final)
                    if numeros:
                        cliente_numero = numeros[0]
                    else:
                        # Tentar extrair número da obra_id
                        obra_id = str(obra_data.get("id", ""))
                        numeros_obra = re.findall(r'\d+', obra_id)
                        if numeros_obra:
                            cliente_numero = numeros_obra[-1]  # Pegar o último número
            
            # Se ainda não tiver número, usar "001" como padrão
            if not cliente_numero:
                cliente_numero = "001"
            # Garantir que o número tenha pelo menos 3 dígitos
            elif len(cliente_numero) < 3:
                cliente_numero = cliente_numero.zfill(3)
            
            # Limpar caracteres não alfanuméricos
            sigla_limpa = re.sub(r'[^a-zA-Z0-9]', '', sigla)
            numero_limpo = re.sub(r'[^a-zA-Z0-9]', '', str(cliente_numero))
            
            return f"{sigla_limpa}_{numero_limpo}"
            
        except Exception as e:
            print(f"❌ Erro ao extrair base do nome: {e}")
            return "EMP_001"
    
    def generate_pc_filename(self, obra_data):
        """Gera nome do arquivo para Proposta Comercial no formato PC_Obra_sigla_numero_data"""
        try:
            from datetime import datetime
            import pytz
            
            try:
                tz = pytz.timezone('America/Sao_Paulo')
                data_atual = datetime.now(tz)
            except:
                data_atual = datetime.now()
            
            # Extrair base do nome (sigla_numero)
            base_name = self.extract_filename_base(obra_data)
            
            # Formatar data como DD-MM-AAAA
            data_formatada = data_atual.strftime("%d-%m-%Y")
            
            # Gerar nome no formato: PC_Obra_sigla_numero_data
            return f"PC_Obra_{base_name}_{data_formatada}.docx"
            
        except Exception as e:
            print(f"❌ Erro ao gerar nome PC: {e}")
            # Fallback: PC_Obra_data
            from datetime import datetime
            data_fallback = datetime.now().strftime("%d-%m-%Y")
            return f"PC_Obra_{data_fallback}.docx"
    
    def generate_pt_filename(self, obra_data):
        """Gera nome do arquivo para Proposta Técnica no formato PT_Obra_sigla_numero_data"""
        try:
            from datetime import datetime
            import pytz
            
            try:
                tz = pytz.timezone('America/Sao_Paulo')
                data_atual = datetime.now(tz)
            except:
                data_atual = datetime.now()
            
            # Extrair base do nome (sigla_numero)
            base_name = self.extract_filename_base(obra_data)
            
            # Formatar data como DD-MM-AAAA
            data_formatada = data_atual.strftime("%d-%m-%Y")
            
            # Gerar nome no formato: PT_Obra_sigla_numero_data
            return f"PT_Obra_{base_name}_{data_formatada}.docx"
            
        except Exception as e:
            print(f"❌ Erro ao gerar nome PT: {e}")
            # Fallback: PT_Obra_data
            from datetime import datetime
            data_fallback = datetime.now().strftime("%d-%m-%YYYY")
            return f"PT_Obra_{data_fallback}.docx"
